"""Document generation service.

Produces formatted CV documents styled to match the AssureSoft reference template:
  - Top-right "Selected by assuresoft" header + full-width rule
  - Large bold candidate name → blue position title
  - Two-column layout: blue section label (left ~21%) | content with top rule (right ~79%)
  - Footer: www.assuresoft.com | page number
  - Sections: Availability, Summary, Experience, Education, Certifications, Skills, Languages

Outputs:
  - DOCX via python-docx
  - PDF  via Jinja2 + WeasyPrint (requires WeasyPrint system dependencies)
"""

import io
import logging
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

from app.core.exceptions import AppException
from app.schemas.document import GenerateCVRequest, LanguageItem

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates"

# ── Brand colours (AssureSoft palette) ───────────────────────────────────────
BLUE  = RGBColor(0x1A, 0x6C, 0xE8)   # section labels, position, bullets
BLACK = RGBColor(0x1A, 0x1A, 0x1A)   # body text, candidate name
GRAY  = RGBColor(0x33, 0x33, 0x33)   # period, italic meta
LGRAY = RGBColor(0xCC, 0xCC, 0xCC)   # divider lines

# ── Column widths (A4 content = 160mm with 25mm margins each side) ────────────
LABEL_COL_MM   = 33   # ~21 %
CONTENT_COL_MM = 127  # ~79 %


# ── Public API ─────────────────────────────────────────────────────────────────

def generate_docx(data: GenerateCVRequest, template_name: str = "assuresoft") -> bytes:
    _validate_template(template_name)
    doc = _build_docx(data, template_name)
    buf = io.BytesIO()
    doc.save(buf)
    logger.info("DOCX generated: candidate=%s template=%s", data.candidate_name, template_name)
    return buf.getvalue()


def generate_pdf(data: GenerateCVRequest, template_name: str = "assuresoft") -> bytes:
    """Render HTML template to PDF via WeasyPrint.

    Requires WeasyPrint + system libraries (libpango, libcairo).
    Works in the Docker image; on Windows install GTK+ runtime first.
    """
    try:
        from jinja2 import Environment, FileSystemLoader
        from weasyprint import HTML
    except ImportError as exc:
        raise AppException(
            "PDF generation requires WeasyPrint. Install: pip install weasyprint",
            status_code=501,
            error_code="pdf_unavailable",
        ) from exc

    _validate_template(template_name)
    template_dir = TEMPLATES_DIR / template_name

    env = Environment(loader=FileSystemLoader(str(template_dir)), autoescape=True)
    template = env.get_template("cv.html")
    html_content = template.render(**_build_template_context(data, template_dir))

    pdf_bytes = HTML(string=html_content, base_url=str(template_dir)).write_pdf()
    logger.info("PDF generated: candidate=%s template=%s", data.candidate_name, template_name)
    return pdf_bytes


# ── Helpers ───────────────────────────────────────────────────────────────────

def _validate_template(name: str) -> None:
    path = TEMPLATES_DIR / name
    if not path.is_dir():
        available = [d.name for d in TEMPLATES_DIR.iterdir() if d.is_dir()]
        raise AppException(
            f"Template '{name}' not found. Available: {', '.join(available)}",
            status_code=400,
            error_code="template_not_found",
        )


def _level_dots(lang: LanguageItem) -> list[int]:
    levels = lang.available_levels or ["Basic", "Intermediate", "Advanced"]
    try:
        idx = levels.index(lang.level)
    except ValueError:
        idx = len(levels) - 1
    return [1 if i <= idx else 0 for i in range(len(levels))]


def _build_template_context(data: GenerateCVRequest, template_dir: Path) -> dict:
    logo = template_dir / "logo-assuresoft.png"
    return {
        "candidate_name": data.candidate_name,
        "position":       data.position,
        "availability":   data.availability,
        "summary":        data.summary,
        "experience":     [e.model_dump() for e in data.experience],
        "education":      [e.model_dump() for e in data.education],
        "certifications": [e.model_dump() for e in data.certifications],
        "skills":         [s.model_dump() for s in data.skills],
        "languages":      [
            {"name": l.name, "level": l.level, "dots": _level_dots(l)}
            for l in data.languages
        ],
        "logo_path": str(logo) if logo.exists() else None,
    }


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _build_docx(data: GenerateCVRequest, template_name: str = "assuresoft") -> Document:
    doc = Document()

    # A4 page, 25mm left/right, 18mm top, 22mm bottom
    for sec in doc.sections:
        sec.page_width    = Mm(210)
        sec.page_height   = Mm(297)
        sec.top_margin    = Mm(18)
        sec.bottom_margin = Mm(22)
        sec.left_margin   = Mm(25)
        sec.right_margin  = Mm(25)

    # Global Normal style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after  = Pt(0)

    _add_page_header(doc, template_name)
    _add_page_footer(doc)
    _add_candidate_block(doc, data)
    _add_section(doc, "Availability",    lambda c: _fill_availability(c, data.availability))
    _add_section(doc, "Summary",         lambda c: _fill_summary(c, data.summary))

    if data.experience:
        _add_section(doc, "Experience",  lambda c: _fill_experience(c, data.experience))

    if data.education:
        _add_section(doc, "Education",   lambda c: _fill_education(c, data.education))

    if data.certifications:
        _add_section(doc, "Certifications", lambda c: _fill_certifications(c, data.certifications))

    if data.skills:
        _add_section(doc, "Skills",      lambda c: _fill_skills(c, data.skills))

    if data.languages:
        _add_section(doc, "Languages",   lambda c: _fill_languages(c, data.languages))

    return doc


# ── Page header & footer ──────────────────────────────────────────────────────

def _add_page_header(doc: Document, template_name: str = "assuresoft") -> None:
    """Top right: logo image (if available) + horizontal rule."""
    header = doc.sections[0].header
    header.is_linked_to_previous = False

    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(4)

    logo_path = TEMPLATES_DIR / template_name / "logo-assuresoft.png"
    if logo_path.exists():
        run = p.add_run()
        run.add_picture(str(logo_path), height=Mm(8))
    else:
        # Fallback text when logo file is not present yet
        r = p.add_run("\u276Fassuresoft")
        r.font.name  = "Arial"
        r.font.size  = Pt(11)
        r.font.bold  = True
        r.font.color.rgb = BLUE

    _add_para_bottom_border(p, color="1A1A1A", size="6")


def _add_page_footer(doc: Document) -> None:
    """Bottom: www.assuresoft.com (left) | page number (right)."""
    footer = doc.sections[0].footer
    footer.is_linked_to_previous = False

    for p in footer.paragraphs:
        p.clear()

    p = footer.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    _add_para_top_border(p, color="1A1A1A", size="6")

    # Tab stop at far right for page number
    p.paragraph_format.tab_stops.add_tab_stop(Mm(160), WD_ALIGN_PARAGRAPH.RIGHT)

    r1 = p.add_run("www.assuresoft.com")
    r1.font.name  = "Arial"
    r1.font.size  = Pt(8.5)
    r1.font.color.rgb = BLUE
    r1.font.underline = True

    p.add_run("\t")

    # Page number field
    r2 = p.add_run()
    r2.font.name  = "Arial"
    r2.font.size  = Pt(8.5)
    r2.font.color.rgb = BLACK
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    r2._r.append(fld)

    r3 = p.add_run()
    instr = OxmlElement("w:instrText")
    instr.text = "PAGE"
    r3._r.append(instr)

    r4 = p.add_run()
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    r4._r.append(fld2)


# ── Candidate name + position ─────────────────────────────────────────────────

def _add_candidate_block(doc: Document, data: GenerateCVRequest) -> None:
    p_name = doc.add_paragraph()
    p_name.paragraph_format.space_before = Pt(6)
    p_name.paragraph_format.space_after  = Pt(2)
    r = p_name.add_run(data.candidate_name)
    r.font.name  = "Arial"
    r.font.size  = Pt(24)
    r.font.bold  = True
    r.font.color.rgb = BLACK

    p_pos = doc.add_paragraph()
    p_pos.paragraph_format.space_after = Pt(14)
    r2 = p_pos.add_run(data.position)
    r2.font.name  = "Arial"
    r2.font.size  = Pt(13)
    r2.font.bold  = True
    r2.font.color.rgb = BLUE


# ── Generic two-column section ────────────────────────────────────────────────

def _add_section(doc: Document, label: str, fill_fn) -> None:
    """Create a two-column borderless table: label | content."""
    table = doc.add_table(rows=1, cols=2)
    _clear_table_borders(table)

    # Column widths
    table.columns[0].width = Mm(LABEL_COL_MM)
    table.columns[1].width = Mm(CONTENT_COL_MM)

    # Left cell — blue label
    left  = table.cell(0, 0)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _clear_cell_borders(left)

    lp = left.paragraphs[0]
    lp.paragraph_format.space_before = Pt(10)
    r = lp.add_run(label)
    r.font.name  = "Arial"
    r.font.size  = Pt(11)
    r.font.bold  = True
    r.font.color.rgb = BLUE

    # Right cell — top rule + content
    right = table.cell(0, 1)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _clear_cell_borders(right)
    _set_cell_top_border(right)

    # Remove the default empty paragraph so fill_fn starts clean
    right._tc.remove(right.paragraphs[0]._p)
    fill_fn(right)

    # Spacer paragraph after the table
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


# ── Section content fillers ───────────────────────────────────────────────────

def _fill_availability(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = BLACK


def _fill_summary(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name  = "Arial"
    r.font.size  = Pt(10)
    r.font.color.rgb = BLACK


def _fill_experience(cell, items) -> None:
    first = True
    for exp in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0) if first else Pt(10)
        p.paragraph_format.space_after  = Pt(0)
        first = False

        r = p.add_run(exp.company)
        r.font.name = "Arial"; r.font.size = Pt(10.5)
        r.font.bold = True; r.font.color.rgb = BLACK

        if exp.position:
            p2 = cell.add_paragraph()
            p2.paragraph_format.space_after = Pt(0)
            r2 = p2.add_run(exp.position)
            r2.font.name = "Arial"; r2.font.size = Pt(10)
            r2.font.bold = True; r2.font.color.rgb = BLACK

        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(4)
        r3 = p3.add_run(exp.period)
        r3.font.name = "Arial"; r3.font.size = Pt(10)
        r3.font.italic = True; r3.font.color.rgb = GRAY

        if exp.summary:
            p4 = cell.add_paragraph()
            p4.paragraph_format.space_after = Pt(3)
            p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r4 = p4.add_run(exp.summary)
            r4.font.name = "Arial"; r4.font.size = Pt(10); r4.font.color.rgb = BLACK

        for ach in exp.achievements:
            _add_bullet(cell, ach)


def _fill_education(cell, items) -> None:
    first = True
    for edu in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0) if first else Pt(6)
        p.paragraph_format.space_after  = Pt(1)
        first = False
        r = p.add_run(edu.degree)
        r.font.name = "Arial"; r.font.size = Pt(10)
        r.font.bold = True; r.font.color.rgb = BLACK

        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(f"{edu.location}  |  {edu.year}")
        r2.font.name = "Arial"; r2.font.size = Pt(10)
        r2.font.italic = True; r2.font.color.rgb = GRAY


def _fill_certifications(cell, items) -> None:
    for cert in items:
        _add_bullet(cell, f"{cert.name}, {cert.institution} ({cert.year})")


def _fill_skills(cell, categories) -> None:
    for cat in categories:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Mm(0)

        bullet = p.add_run("●  ")
        bullet.font.name = "Arial"; bullet.font.size = Pt(8)
        bullet.font.bold = True; bullet.font.color.rgb = BLUE

        label = p.add_run(f"{cat.area}:")
        label.font.name = "Arial"; label.font.size = Pt(10)
        label.font.bold = True; label.font.color.rgb = BLACK

        for item in cat.items:
            pi = cell.add_paragraph()
            pi.paragraph_format.left_indent = Mm(5)
            pi.paragraph_format.space_after = Pt(1)

            sub = pi.add_run("○  ")
            sub.font.name = "Arial"; sub.font.size = Pt(8)
            sub.font.color.rgb = GRAY

            ri = pi.add_run(item)
            ri.font.name = "Arial"; ri.font.size = Pt(10); ri.font.color.rgb = BLACK


def _fill_languages(cell, items) -> None:
    for lang in items:
        _add_bullet(cell, f"{lang.name} - {lang.level}")


# ── Low-level XML helpers ─────────────────────────────────────────────────────

def _add_bullet(cell, text: str) -> None:
    p = cell.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Mm(0)

    dot = p.add_run("●  ")
    dot.font.name = "Arial"; dot.font.size = Pt(8)
    dot.font.bold = True; dot.font.color.rgb = BLUE

    r = p.add_run(text)
    r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = BLACK


def _clear_table_borders(table) -> None:
    tbl  = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    bdr  = OxmlElement("w:tblBorders")
    for side in ("top", "start", "bottom", "end", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"),  "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        bdr.append(el)
    tblPr.append(bdr)


def _clear_cell_borders(cell) -> None:
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    bdr   = OxmlElement("w:tcBorders")
    for side in ("top", "start", "bottom", "end"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bdr.append(el)
    tcPr.append(bdr)


def _set_cell_top_border(cell, color: str = "CCCCCC", size: str = "4") -> None:
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    bdr  = OxmlElement("w:tcBorders")

    top = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    size)
    top.set(qn("w:space"), "0")
    top.set(qn("w:color"), color)
    bdr.append(top)

    for side in ("start", "bottom", "end"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        bdr.append(el)

    tcPr.append(bdr)


def _add_para_bottom_border(p, color: str = "1A1A1A", size: str = "6") -> None:
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    size)
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_para_top_border(p, color: str = "1A1A1A", size: str = "6") -> None:
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top  = OxmlElement("w:top")
    top.set(qn("w:val"),   "single")
    top.set(qn("w:sz"),    size)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)
