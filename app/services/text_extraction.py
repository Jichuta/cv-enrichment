"""Text extraction service — pulls plain text from PDF and DOCX CV files.

Libraries used:
  - pdfplumber: best-in-class PDF text extraction; handles multi-column
    layouts, tables, and embedded fonts better than pypdf/PyMuPDF for
    structured documents like CVs.
  - python-docx: the standard library for reading .docx files; traverses
    paragraphs and table cells to preserve all text.
"""

import io
import logging

import pdfplumber
from docx import Document

from app.core.exceptions import ExtractionError
from app.schemas.extraction import TextExtractionResponse

logger = logging.getLogger(__name__)

# Allowed MIME types and their canonical short names
SUPPORTED_TYPES: dict[str, str] = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}

MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024


def extract_text(file_bytes: bytes, filename: str, content_type: str) -> TextExtractionResponse:
    """Extract plain text from a PDF or DOCX file.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used in the response and logs).
        content_type: MIME type declared by the client.

    Returns:
        TextExtractionResponse with the raw text and document metadata.

    Raises:
        ExtractionError: If the file is corrupt or text cannot be extracted.
    """
    file_type = SUPPORTED_TYPES[content_type]  # caller must validate before calling

    try:
        if file_type == "pdf":
            return _extract_pdf(file_bytes, filename)
        return _extract_docx(file_bytes, filename)
    except ExtractionError:
        raise
    except Exception as exc:
        logger.exception("Unexpected error extracting text from '%s'", filename)
        raise ExtractionError(
            f"Failed to extract text from '{filename}'.",
            details={"reason": str(exc)},
        ) from exc


def _extract_pdf(file_bytes: bytes, filename: str) -> TextExtractionResponse:
    pages: list[str] = []

    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        if not pdf.pages:
            raise ExtractionError(
                f"'{filename}' appears to be an empty PDF with no pages."
            )

        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=3)
            pages.append(text or "")

    raw_text = "\n\n".join(p for p in pages if p.strip())

    if not raw_text.strip():
        raise ExtractionError(
            f"No readable text found in '{filename}'. "
            "The PDF may be image-based (scanned) and requires OCR."
        )

    logger.info("PDF extracted: file=%s pages=%d chars=%d", filename, len(pages), len(raw_text))
    return _build_response(filename, "pdf", raw_text, len(pages))


def _extract_docx(file_bytes: bytes, filename: str) -> TextExtractionResponse:
    doc = Document(io.BytesIO(file_bytes))

    chunks: list[str] = []

    # Paragraphs (main body text, headings, list items)
    for para in doc.paragraphs:
        if para.text.strip():
            chunks.append(para.text.strip())

    # Tables — common in CV layouts for skills grids, etc.
    for table in doc.tables:
        for row in table.rows:
            row_text = "  ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                chunks.append(row_text)

    raw_text = "\n".join(chunks)

    if not raw_text.strip():
        raise ExtractionError(f"No readable text found in '{filename}'.")

    # DOCX has no concept of pages — estimate based on word count (~250 words/page)
    word_count = len(raw_text.split())
    estimated_pages = max(1, round(word_count / 250))

    logger.info(
        "DOCX extracted: file=%s estimated_pages=%d chars=%d",
        filename,
        estimated_pages,
        len(raw_text),
    )
    return _build_response(filename, "docx", raw_text, estimated_pages)


def _build_response(
    filename: str, file_type: str, raw_text: str, page_count: int
) -> TextExtractionResponse:
    return TextExtractionResponse(
        filename=filename,
        file_type=file_type,
        raw_text=raw_text,
        page_count=page_count,
        word_count=len(raw_text.split()),
        char_count=len(raw_text),
    )
